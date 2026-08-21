from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn
from docx.shared import Inches

PATH = 'Ashfall_DM_Campaign_Guide_v5.docx'

def insert_before(paragraph, new_paragraph):
    paragraph._p.addprevious(new_paragraph._p)

def add_para_before(anchor, text='', style='Normal'):
    p = OxmlElement('w:p')
    anchor._p.addprevious(p)
    para = Paragraph(p, anchor._parent)
    para.style = style
    if text:
        para.add_run(text)
    return para

def add_table_before(anchor, rows):
    tbl = anchor._parent.add_table(rows=0, cols=len(rows[0]), width=Inches(6.5))
    anchor._p.addprevious(tbl._tbl)
    for row_i, values in enumerate(rows):
        cells = tbl.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
    return tbl

doc = Document(PATH)
anchor = next(p for p in doc.paragraphs if p.style.name == 'Heading 1' and p.text == 'The true story of Ashfall')

add_para_before(anchor, 'First encounters on the road', 'Heading 1')
add_para_before(anchor, 'Use these two short scenes before the party reaches Ashfall. They are intentionally self-contained: the combat teaches movement, target choice and mercy; the roadside problem teaches investigation, cooperation and consequence. Neither encounter requires knowledge of the village, the crystal or the loop.', 'Normal')

add_para_before(anchor, 'Encounter 1 — Ash on the Road', 'Heading 2')
add_para_before(anchor, 'Purpose: a low-stakes first fight with clear enemy behaviour and a choice that rewards observation. Run it for 3–5 characters at level 1; adjust by removing the swarm for a very small or fragile party.', 'Normal')
add_para_before(anchor, 'Read aloud', 'Heading 3')
add_para_before(anchor, 'The road dips between black-barked trees. A grey animal steps out of the ash ahead, its paws making no sound. A second shape circles behind you. When the wind shifts, the creatures’ ribs glow like banked coals—and something small and bright skitters through the dust at your feet.', 'Normal')
add_table_before(anchor, [
    ['Opposition', 'Simple running notes'],
    ['2 ash-hounds', 'AC 12; 7 HP each; bite +4, 1d6+2 piercing. They focus on whoever is isolated, then retreat if both are bloodied.'],
    ['1 cinder swarm', 'AC 11; 10 HP; ember spray +3, 1d4 fire to one nearby creature. It occupies a 5-foot space and flees when reduced below 4 HP.'],
])
add_para_before(anchor, 'Three-beat flow', 'Heading 3')
add_para_before(anchor, '1. The hounds test the edges of the group rather than charging the strongest character. 2. The swarm crawls toward dropped gear, a waterskin or a campfire, giving the players an obvious environmental problem. 3. If a character uses water, a cloak, sand or a convincing threat to break the pack’s nerve, the creatures flee instead of fighting to the end.', 'Normal')
add_para_before(anchor, 'DM cues and outcomes', 'Heading 3')
add_table_before(anchor, [
    ['If the party…', 'Then…'],
    ['Uses a smart formation or protects a vulnerable ally', 'Let the hounds lose interest in that target and reward the tactic with a clear opening.'],
    ['Investigates the ash or watches the swarm', 'They notice the creatures avoid living green plants and can gain advantage on the next attempt to drive them off.'],
    ['Wins by force', 'The road is safe. One hound leaves a warm, glassy tooth that can be sold or shown later as a mundane omen.'],
    ['Wins by mercy or deterrence', 'The creatures retreat into the trees. Give the party a small Bond-leaning moment: they solved a danger without needless killing.'],
])
add_para_before(anchor, 'Keep it spoiler-light: the ash is ordinary magical fallout from a dry valley, not a clue to the Unremembered or the origin of Ashfall.', 'Normal')

add_para_before(anchor, 'Encounter 2 — The Empty Cart', 'Heading 2')
add_para_before(anchor, 'Purpose: a compact non-combat scene that lets every character contribute. The party finds a stranded peddler, Orin Pell, beside a cart that appears to have lost the same wheel three times. Nothing supernatural is happening; exhaustion, a bent axle and a frightened mule are enough.', 'Normal')
add_para_before(anchor, 'Read aloud', 'Heading 3')
add_para_before(anchor, 'A handcart sits crooked across the road. Its owner is kneeling beside one wheel, staring at the empty space where a wooden pin should be. A mule trembles in the traces. “I had it a moment ago,” the man says. “I keep putting it down, and then it is gone.” Around him, three identical bundles of rope lie in the dust.', 'Normal')
add_table_before(anchor, [
    ['Approach', 'Useful check / DC', 'What it reveals'],
    ['Calm Orin and the mule', 'Animal Handling or Persuasion 11', 'Orin can explain the route and admits he has not slept since yesterday.'],
    ['Inspect the cart', 'Investigation or Perception 12', 'The pin is caught in a torn blanket under the cart, not missing.'],
    ['Repair or improvise', 'Sleight of Hand, Survival or carpenter’s tools 12', 'A tent peg, dagger hilt or tied cord can hold the wheel until the next village.'],
    ['Coordinate the group', 'Any clever plan; no roll if everyone describes a useful role', 'The party solves the problem quickly and Orin offers directions and a small ration bundle.'],
])
add_para_before(anchor, 'Scene rhythm', 'Heading 3')
add_para_before(anchor, 'Start with Orin talking too quickly. After one useful question, let a player notice the pin under the blanket. If the group argues or fails a check, the mule pulls free and the cart rolls a few feet downhill; the problem becomes urgent, not hopeless. Once the cart is safe, ask each player what small task their character chose.', 'Normal')
add_para_before(anchor, 'Outcomes', 'Heading 3')
add_table_before(anchor, [
    ['Result', 'Reward'],
    ['The party helps without taking over', 'Orin shares dried fruit, gives directions and remembers the party as competent travellers.'],
    ['The party rushes or frightens him', 'The cart is repaired, but Orin leaves wary; the party learns that success can still cost trust.'],
    ['The party abandons him', 'They can continue, but the next stretch is harder: no directions, no ration bundle and a reminder that choices have social weight.'],
])
add_para_before(anchor, 'Optional connective tissue: Orin has heard that the road into the valley is safe during daylight and that Ashfall has a good inn. He knows nothing about the village’s history and should not mention loops, crystals, witches or the Unremembered.', 'Normal')

doc.save(PATH)
print('Added first encounters section')
