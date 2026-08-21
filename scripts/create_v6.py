from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "Ashfall_DM_Campaign_Guide_v5.docx"
OUTPUT = ROOT / "Ashfall_DM_Campaign_Guide_v6.docx"


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table: Table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row_index, row in enumerate(table.rows):
        if row_index == 0:
            tbl_header = OxmlElement("w:tblHeader")
            tbl_header.set(qn("w:val"), "true")
            row._tr.get_or_add_trPr().append(tbl_header)
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def table_widths(column_count: int, key: str | None = None) -> list[int]:
    if key == "memory":
        return [1800, 1800, 5760]
    if key == "timeline":
        return [1500, 2560, 5300]
    if key == "clues":
        return [2850, 6510]
    if column_count == 2:
        return [3000, 6360]
    if column_count == 3:
        return [2100, 3000, 4260]
    return [9360 // column_count] * column_count


def add_paragraph_before(document: Document, marker, text: str, style: str = "Normal") -> Paragraph:
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(text)
    marker.addprevious(paragraph._p)
    if style.startswith("Heading"):
        paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_table_before(document: Document, marker, rows: list[list[str]], key: str | None = None) -> Table:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    for row_index, values in enumerate(rows):
        for cell, value in zip(table.rows[row_index].cells, values):
            set_cell_text(cell, value, bold=row_index == 0)
            if row_index == 0:
                shade_cell(cell, "DCE6F1")
    set_table_geometry(table, table_widths(len(rows[0]), key))
    marker.addprevious(table._tbl)
    return table


def add_blocks(document: Document, marker, blocks: list[dict]) -> None:
    for block in blocks:
        block_type = block["type"]
        if block_type == "p":
            add_paragraph_before(document, marker, block["text"], block.get("style", "Normal"))
        elif block_type == "h2":
            add_paragraph_before(document, marker, block["text"], "Heading 2")
        elif block_type == "h3":
            add_paragraph_before(document, marker, block["text"], "Heading 3")
        elif block_type == "list":
            for item in block["items"]:
                add_paragraph_before(document, marker, item, "List Bullet")
        elif block_type == "table":
            add_table_before(document, marker, block["rows"], block.get("key"))
        elif block_type == "callout":
            add_table_before(document, marker, [[block["text"]]], block.get("key"))
        else:
            raise ValueError(f"Unsupported block type: {block_type}")


def heading_paragraph(document: Document, title: str) -> Paragraph:
    return next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.style.name == "Heading 1" and paragraph.text.strip() == title
    )


def next_heading_marker(document: Document, heading: Paragraph):
    body = document.element.body
    children = list(body)
    start = children.index(heading._p)
    for child in children[start + 1 :]:
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, document)
            if paragraph.style.name == "Heading 1":
                return child
    return body.sectPr


def replace_section(document: Document, title: str, blocks: list[dict]) -> None:
    heading = heading_paragraph(document, title)
    marker = next_heading_marker(document, heading)
    body = document.element.body
    children = list(body)
    start = children.index(heading._p)
    end = children.index(marker)
    for child in children[start + 1 : end]:
        body.remove(child)
    add_blocks(document, marker, blocks)


def insert_section_before(document: Document, before_title: str, title: str, blocks: list[dict]) -> None:
    marker = heading_paragraph(document, before_title)._p
    add_paragraph_before(document, marker, title, "Heading 1")
    add_blocks(document, marker, blocks)


def replace_text_in_paragraph(paragraph: Paragraph, old: str, new: str) -> bool:
    if old not in paragraph.text:
        return False
    full = paragraph.text.replace(old, new)
    paragraph.text = full
    return True


def replace_text_everywhere(document: Document, old: str, new: str) -> int:
    count = 0
    for paragraph in document.paragraphs:
        count += int(replace_text_in_paragraph(paragraph, old, new))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    count += int(replace_text_in_paragraph(paragraph, old, new))
    return count


def main() -> None:
    shutil.copyfile(SOURCE, OUTPUT)
    document = Document(OUTPUT)
    document.core_properties.title = "Ashfall: The Day That Refuses to Die - DM Campaign Guide v6"
    document.core_properties.subject = "D&D time-loop mystery campaign guide, first-loop runner and NPC bible"

    replace_section(
        document,
        "Quick start",
        [
            {
                "type": "callout",
                "text": "Use this guide, not every detail. For the first session prepare Elias, Mara, Nessa, the Lantern House fire and the 3:13 a.m. reset. The first loop establishes a normal Ashfall; the second loop begins the investigation.",
            },
            {"type": "h2", "text": "Campaign at a glance"},
            {
                "type": "table",
                "rows": [
                    ["Question", "Campaign answer"],
                    ["What is Ashfall?", "A magical reconstruction of a village destroyed 307 years ago. Its repeated identities form a prison around the Unremembered; Seraphine's dead echo maintains the seal."],
                    ["Why do the heroes remember?", "The Witness Crystal anchors living outsiders across resets."],
                    ["What is Elias?", "The Unremembered assembled Elias from stolen village memories as an unwitting escape plan. Once the crystal anchored him, he became an independent person the entity cannot control."],
                    ["Why does Elias need the party?", "The final ward detects the Unremembered's signature in his creation and rejects him. Only independent living outsiders can authorise a fundamental change."],
                    ["What does Elias believe?", "He remembers being an Ashfall villager and believes termination will return his people to ordinary life. His love is genuine; his origin and expected outcome are false."],
                    ["What drives play?", "Care about the villagers, witness the reset, test what persists, follow character-led clues and decide when enough is known to descend beneath the church."],
                    ["What changes the ending?", "Bond, Truth and Fracture imprints stored in the crystal. Number of loops alone changes nothing."],
                ],
                "key": "clues",
            },
            {"type": "h2", "text": "The memory rule"},
            {
                "type": "callout",
                "text": "Elias remembers the plan. Nessa remembers the people. Seraphine remembers the disaster. Everyone else remembers only through evidence and emotion.",
            },
            {
                "type": "table",
                "rows": [
                    ["Character", "Awareness", "What survives a reset"],
                    ["Player characters", "Full", "Their memories, identities and outside equipment are anchored by the Witness Crystal."],
                    ["Elias", "Full", "The crystal permanently anchors his constructed identity. The Unremembered cannot rewrite or command him."],
                    ["Nessa", "Partial", "People, promises and powerful moments remain. She does not retain a perfect timeline or understand the prison."],
                    ["Seraphine", "Fragmented", "Her echo perceives the resets through the spell but can communicate only through supported anchors."],
                    ["Mara, Aldren, Iven, Silas and other villagers", "None", "Facts reset. Strong trust, fear, guilt, dreams and habits may carry over emotionally."],
                    ["Preserved records", "Evidence", "Outside or crystal-anchored writing can survive even when its author forgets creating it."],
                ],
                "key": "memory",
            },
            {"type": "h2", "text": "What to prepare before the first session"},
            {
                "type": "list",
                "items": [
                    "Assume the party reaches the Lantern House at about 1 p.m.",
                    "Prepare only Elias, Mara and Nessa as active characters. Aldren or Iven may appear briefly at the festival or fire.",
                    "Use the Lantern House fire as the first meaningful crisis. Record one person, promise or object the party chose to protect.",
                    "At 3:12 a.m. the crystal wakes every anchored character automatically. Never hide the first reset behind a Perception check.",
                    "After the reset, show the inn restored, the harmed villagers alive, ordinary villagers forgetful and Elias clearly unsurprised.",
                    "Put three bowls, dice or marks on paper for Bond, Truth and Fracture.",
                ],
            },
            {"type": "h2", "text": "The central rule"},
            {
                "type": "callout",
                "text": "Loops do not automatically damage the spell. A reset restores Ashfall's original pattern. Only changes preserved by the Witness Crystal matter to the finale. Reckless destruction without the crystal can erase Ashfall, but it cannot safely open or repair the prison.",
            },
        ],
    )

    replace_section(
        document,
        "The true story of Ashfall",
        [
            {"type": "h2", "text": "What happened 307 years ago"},
            {
                "type": "list",
                "items": [
                    "Miners beneath Saint Orra's church breached an ancient chamber and awakened the Unremembered, an entity that consumes identity and uses stolen memories as disguises.",
                    "Seraphine Vey, Ashfall's resident archmage and protector, warned the mayor and tried to seal the excavation. Most villagers distrusted her secrecy and blamed her for the ash, nightmares and disappearances.",
                    "When the entity emerged, the catastrophe killed the villagers. Seraphine used their final memories to reconstruct Ashfall's last day as a closed repeating pattern before the entity could carry their identities beyond the valley.",
                    "The reconstructed villagers became the walls of the prison. Their thoughts, feelings and relationships are genuine even though their original bodies lie beneath the cemetery.",
                    "Seraphine's body died completing the spell. A damaged echo of her consciousness remains distributed through its anchors, maintaining the seal and supporting the Witness Crystal as the authorised interface for future living outsiders.",
                    "As the prison closed, the Unremembered used stolen memories to assemble Elias Voss. It gave him a false life in Ashfall and a false memory that Seraphine sent him for help. The crystal recognised a distinct mind and anchored him, placing Elias beyond the entity's direct control.",
                    "The entity cannot command Elias. It gambled that a sincere villager who loved Ashfall would eventually persuade independent outsiders to terminate the spell, unravel the villagers and open the prison.",
                ],
            },
            {"type": "h2", "text": "What the party initially believes"},
            {
                "type": "list",
                "items": [
                    "Elias is a travelling merchant and former Ashfall villager transporting a sealed magical relic.",
                    "Ashfall is isolated but inhabited and preparing for a harvest festival.",
                    "The burned woman seen in mirrors may be the witch responsible for the village's recent omens.",
                    "Elias escaped a curse and has returned with the crystal so outsiders can restore his village to ordinary time.",
                ],
            },
            {"type": "h2", "text": "What they should eventually learn"},
            {
                "type": "list",
                "items": [
                    "The villagers died centuries ago, but their reconstructed minds are genuine people.",
                    "Seraphine caused neither the catastrophe nor the villagers' deaths. She is the dead protector maintaining the entity's prison, although she preserved the villagers without their consent.",
                    "Elias never lived in the original Ashfall. His memories were assembled from stolen fragments belonging to several villagers.",
                    "The Unremembered created Elias but cannot control him. His love for Ashfall and every choice he made after being anchored are real.",
                    "Elias has brought earlier groups and sincerely believes termination will restore the villagers. Instead it will unravel them and release the Unremembered.",
                    "The crystal is learning from the party's choices. It can preserve Ashfall, rewrite it or invoke termination, but there is no consequence-free ending.",
                ],
            },
        ],
    )

    replace_section(
        document,
        "Elias and the crystal",
        [
            {
                "type": "callout",
                "text": "The Unremembered created Elias's beginning, but it does not own the person he became. Elias is neither possessed nor secretly evil. He is an independent person acting from genuine love on the basis of manufactured memories.",
            },
            {"type": "h2", "text": "What Elias actually is"},
            {
                "type": "p",
                "text": "As Seraphine's prison closed, the Unremembered assembled a human-shaped identity from memories it had stolen from Ashfall's dead. That identity was Elias Voss. The entity could not cross the seal itself, so it created someone who would want to return with outside witnesses and open it lawfully.",
            },
            {
                "type": "p",
                "text": "The Witness Crystal recognised Elias as a distinct mind and permanently anchored him. From that moment his identity became continuous and protected. The Unremembered cannot command him, replace his memories or speak through him. Three centuries of independent choices have made Elias far more than the purpose behind his creation.",
            },
            {"type": "h2", "text": "The life he remembers"},
            {
                "type": "list",
                "items": [
                    "Growing up in Ashfall and being fed at Mara's table.",
                    "Working as a village courier for Aldren and occasionally carrying messages for Seraphine.",
                    "Fleeing the catastrophe with the Witness Crystal.",
                    "Seraphine telling him to find independent outsiders who could save Ashfall.",
                ],
            },
            {
                "type": "p",
                "text": "Every memory is emotionally convincing, but the details belong to different villagers. There was no original Elias Voss. Mara may recognise Tomas's childhood in one story; Aldren remembers another courier performing Elias's errands; the oldest records contain no birth, home or family bearing his name.",
            },
            {"type": "h2", "text": "What he wants"},
            {
                "type": "p",
                "text": "Elias believes the repeating spell prevents his people from returning to ordinary time. He expects lawful termination to stop the reset while leaving Ashfall and its inhabitants alive, free to remember, age and build a future. He wants to save his home, not destroy it.",
            },
            {"type": "h2", "text": "Why he needs the party"},
            {
                "type": "list",
                "items": [
                    "The inner wards detect the Unremembered's signature in the magic that formed him and reject him as a witness.",
                    "Elias cannot enter the final decision circle or perceive its deepest warning: the villagers are the lock and termination opens the prison.",
                    "The ritual requires several independent living wills so no part of Ashfall, no single rescuer and no compromised victim can rewrite the village alone.",
                    "Only outsiders can create genuinely new memories that did not originate inside the preserved day or the entity's stolen archive.",
                ],
            },
            {"type": "h2", "text": "Why he uses an escort story"},
            {
                "type": "p",
                "text": "Elias has tried telling recruits about the loop before entering Ashfall. Some refused; others treated the villagers as disposable copies. He now poses as a merchant escorting a sealed relic, lets the party meet his people and allows the first reset to prove the impossible. He lies deliberately, but he believes the deception is necessary to give Ashfall a fair chance.",
            },
            {"type": "h2", "text": "His first-loop plan"},
            {
                "type": "list",
                "items": [
                    "Reach the Lantern House around 1 p.m. and claim the chest cannot be delivered until morning.",
                    "Keep the party in Ashfall long enough to meet Mara, Nessa and the festival crowd.",
                    "Protect the chest during the evening crisis, even if that makes his priorities look suspicious.",
                    "Wait for the crystal to wake the party at 3:12 a.m. and let them witness the reset.",
                    "Afterward admit that Ashfall repeats and the crystal preserves memory, but conceal the earlier expeditions and his uncertainty about the final ward.",
                ],
            },
            {"type": "h2", "text": "The tragedy of his plan"},
            {
                "type": "p",
                "text": "Lawful termination does not return Ashfall to ordinary life. It unravels the reconstructed villagers whose identities form the prison and gives the Unremembered a route into the wider world. The entity planted no command in Elias; it constructed a persuasive false past and trusted his compassion to carry the plan forward.",
            },
            {"type": "h2", "text": "The question at the centre of Elias"},
            {
                "type": "callout",
                "text": "If my memories were manufactured, is my love still real? Yes. His origin is a lie, but the person who lived with that origin for three centuries is real and can choose differently.",
            },
        ],
    )

    replace_section(
        document,
        "Player-paced story structure",
        [
            {"type": "p", "text": "Do not structure the campaign by loop number. Use five story states. The first reset moves the party from Arrival to Recognition; after that, characters and clues move them forward at their own pace."},
            {
                "type": "table",
                "rows": [
                    ["Story state", "Your job as DM"],
                    ["1. Arrival", "Begin near 1 p.m. Make Ashfall welcoming, introduce Elias, Mara and Nessa, run the festival and one crisis, then wake the party for the 3:13 a.m. reset."],
                    ["2. Recognition", "Use restored fire damage, living victims, forgotten conversations and Elias's calm reaction to prove the reset beyond doubt."],
                    ["3. Investigation", "Offer NPC agendas and overlapping clues: Iven tests evidence, Aldren creates pressure, Silas explains mechanisms and Nessa remembers promises."],
                    ["4. Revelation", "Confirm the villagers are reconstructions, Seraphine died protecting the region, Elias never lived in Ashfall and termination would release the Unremembered."],
                    ["5. Decision", "Open the church threshold. Let the party descend when they choose and make the finale reflect what the crystal learned."],
                ],
                "key": "clues",
            },
            {"type": "h2", "text": "What must happen eventually"},
            {
                "type": "list",
                "items": [
                    "The party experiences at least one reset and understands that the crystal preserves their memories.",
                    "They discover Elias has brought earlier groups and that his memories do not match Ashfall's records.",
                    "They receive credible evidence that the villagers died 307 years ago and now form part of the prison.",
                    "They learn the Unremembered created Elias but does not control him.",
                    "They receive at least one reliable warning that termination without replacement containment opens the prison.",
                    "They decide when to enter the chamber and which ending to attempt.",
                ],
            },
            {"type": "h2", "text": "When the party stalls"},
            {
                "type": "list",
                "items": [
                    "Advance a scheduled event: the inn fire, Aldren's detention, the festival or the bell.",
                    "Have Nessa recall one promise or feeling and ask what the party will do differently.",
                    "Let Iven propose a fact that can be tested after the next reset.",
                    "Give Elias a helpful action followed by one suspicious omission.",
                    "Move the needed clue to the character or location the party chose to investigate.",
                ],
            },
        ],
    )

    insert_section_before(
        document,
        "Read-aloud descriptions",
        "Running the first loop",
        [
            {
                "type": "callout",
                "text": "The first loop is a baseline, not an investigation. Introduce three people, create one meaningful consequence and then visibly erase it. Do not make the players fill fourteen hours or stay awake by choice.",
            },
            {"type": "h2", "text": "Default first-loop timeline"},
            {
                "type": "table",
                "rows": [
                    ["Time", "Scene", "What it accomplishes"],
                    ["1:00 p.m.", "Arrival at the Lantern House", "Elias says the Crystal Chest cannot be delivered until morning. Mara offers rooms. Nessa notices that the chest makes her ward token feel warm."],
                    ["2:00-5:00 p.m.", "A little ordinary Ashfall", "Allow one location or conversation. Use only subtle details: Tomas's untouched place, a wrong clock or Elias insisting they remain overnight."],
                    ["5:00-8:00 p.m.", "Harvest festival", "Make the village worth caring about. Let the party meet neighbours and make one small personal connection."],
                    ["Around 9:00 p.m.", "Lantern House fire", "Give the party a clear crisis: rescue someone, protect Tomas's belongings, investigate the cause or notice Elias protecting the chest first."],
                    ["After the crisis", "Time cut to sleep", "Resolve a brief aftermath, record the party's strongest choice and allow them to rest. Do not play every remaining hour."],
                    ["3:12 a.m.", "The crystal wakes them", "All anchored characters wake automatically. The chest hums, the bell strikes thirteen times and ash rises instead of falling."],
                    ["3:13 a.m.", "Reset", "The village returns to 6:13 a.m. around the conscious party: fire damage reverses, objects return and the dead or injured are restored."],
                    ["6:13 a.m.", "Undeniable proof", "Mara brings breakfast without remembering them, the inn is whole, Nessa retains a promise and Elias is clearly unsurprised."],
                ],
                "key": "timeline",
            },
            {"type": "h2", "text": "If the party sleeps early"},
            {
                "type": "list",
                "items": [
                    "If they sleep before the festival, the fire can wake them.",
                    "If they prevent the fire, the thirteen bells and crystal wake them.",
                    "Even magical silence cannot suppress the crystal's direct warning to an anchored identity.",
                    "No Perception check is required. The first reset is a campaign event, not a clue they can accidentally miss.",
                ],
            },
            {"type": "h2", "text": "Reset positioning and resources"},
            {
                "type": "list",
                "items": [
                    "Ashfall returns to 6:13 a.m.; it does not return to the party's 1 p.m. arrival.",
                    "The party remains at its current location while the reconstructed village changes around them. Move a character only if the restored location would place them inside solid matter or immediate danger.",
                    "Equipment brought from outside remains with the party. Ashfall-made objects return to their original positions unless the crystal explicitly anchors them.",
                    "For a short campaign, grant the benefit of a long rest at the reset. This keeps bookkeeping simple and makes experimentation less punishing.",
                ],
            },
            {"type": "h2", "text": "First reset - read aloud"},
            {
                "type": "callout",
                "text": "You wake at the same instant. The Crystal Chest is humming beneath the floorboards. Outside, the church bell begins to ring. By the thirteenth stroke, ash is rising past the window instead of falling. The burned walls groan. Black timber becomes pale wood. Smoke pours backwards into the chimney. Broken glass leaps from the floor and seals itself into the window. Darkness snaps into grey morning. Elias is already standing beside the chest.",
            },
        ],
    )

    replace_section(
        document,
        "The repeating day",
        [
            {"type": "p", "text": "Use the schedule as predictable opportunities, not a railroad. The party's first loop begins around 1 p.m.; every later loop begins when Ashfall resets to 6:13 a.m."},
            {
                "type": "table",
                "rows": [
                    ["Time", "Events available"],
                    ["6:13 a.m. - morning", "Mara prepares breakfast and Tomas's untouched place; Aldren prepares the festival announcement; Nessa wakes with emotional fragments; Elias checks the chest; Silas briefly leaves his shop."],
                    ["1:00 p.m. - afternoon", "The party first arrives. Iven patrols; the manor records are accessible; the church is lightly guarded; the cemetery remains blank."],
                    ["5:00 p.m. - evening", "The festival begins; the Lantern House fire may start; Elias protects or prepares the crystal; Aldren's recurring fear focuses on a supposed threat."],
                    ["Midnight - 3:12 a.m.", "Aldren may detain or endanger a chosen villager; names appear in the cemetery; Seraphine can manifest through an anchor; the church stair can be opened."],
                    ["3:13 a.m.", "The crystal wakes anchored minds and Ashfall resets to 6:13 a.m."],
                ],
                "key": "clues",
            },
            {"type": "h2", "text": "Reset procedure"},
            {
                "type": "list",
                "items": [
                    "At 3:12 a.m., wake every crystal-anchored character automatically.",
                    "Ask each player for one image, sound or emotion their character remembers most strongly from the day.",
                    "Restore Ashfall's original physical state. Keep outside equipment and explicitly crystal-anchored objects where they are.",
                    "Keep the party in its current safe position and return the village clock to 6:13 a.m.",
                    "Keep the party's memories. For the simplest short campaign, also restore hit points, spell slots and other long-rest resources.",
                    "Give one or two villagers emotional carryover based on the party's strongest actions. Do not give them factual memory unless a specific exception says otherwise.",
                    "Show one crystal anomaly only when a major imprint was added. Another loop by itself causes no deterioration.",
                ],
            },
            {"type": "h2", "text": "Emotional carryover"},
            {
                "type": "table",
                "rows": [
                    ["Previous treatment", "Next-loop reaction"],
                    ["Protected or comforted", "Unexplained trust, an offered key, food saved for them or a warning given privately."],
                    ["Threatened or killed", "Flinching, avoidance, hostility, nightmares or instinctive recognition of a weapon."],
                    ["Revealed a painful truth", "Grief without context, repeated phrases, dreams or a compulsion to investigate."],
                    ["Inspired a new choice", "The villager hesitates before repeating the normal routine and may eventually act differently."],
                    ["Proved the loop to Aldren", "He forgets the proof after reset but carries unease or guilt. He never begins a loop already knowing that his victim will return."],
                ],
                "key": "clues",
            },
        ],
    )

    replace_section(
        document,
        "Clue ladder",
        [
            {"type": "callout", "text": "Three-clue rule: for every conclusion the party must reach, provide at least three independent clues. A failed roll changes cost, timing or danger, not whether the only clue exists."},
            {
                "type": "table",
                "rows": [
                    ["Conclusion", "Possible clues"],
                    ["The day repeats", "The restored Lantern House; an endangered villager alive at breakfast; an outside object remains moved; Nessa remembers a promise."],
                    ["The crystal preserves memory", "It pulses before the reset; Silas's diagram calls it a witness anchor; distance from it blurs a recent memory; Elias admits the connection."],
                    ["The villagers died centuries ago", "Gravestones; bones beneath the cemetery; ledgers ending 307 years ago; Seraphine's memory of the catastrophe."],
                    ["Elias has done this before", "Portraits with earlier parties; Iven's surviving notes; Mara's dreams; an earlier witness echo recognises his escort story."],
                    ["Elias never lived in Ashfall", "No birth or household record; Mara recognises Tomas's childhood in his stories; Aldren remembers another courier; Seraphine says she never sent him."],
                    ["The Unremembered created Elias but cannot control him", "His memories combine several stolen lives; the inner ward marks his origin with the entity's signature; the crystal protects his stable independent identity; he acts against the entity after learning the truth."],
                    ["Seraphine is protecting the world", "Her warnings target termination; Silas confesses she died completing the seal; the core wards name her Keeper; old records show the village blamed her before the catastrophe."],
                    ["The villagers form the prison", "Church mosaics show linked names around a dark void; Silas explains identity as the seal; removing a name weakens a ward; Seraphine confirms it at a core anchor."],
                    ["Termination opens the prison", "The inner ward reads 'No release without a second vessel'; Silas's diagrams show village and entity on the same lock; an earlier witness says 'the door opens both ways'; a genuine Seraphine manifestation shows a name being removed."],
                ],
                "key": "clues",
            },
            {"type": "h2", "text": "Never hide these behind a roll"},
            {
                "type": "list",
                "items": [
                    "That the reset occurred and the crystal is connected to the party's memories.",
                    "That Elias lied about the escort job and has brought earlier groups.",
                    "That Elias's memories conflict with Ashfall's records.",
                    "That the final choice will change or end Ashfall.",
                    "At least one reliable warning that termination without replacement containment opens the prison.",
                ],
            },
        ],
    )

    replace_section(
        document,
        "NPC quick reference",
        [
            {
                "type": "table",
                "rows": [
                    ["NPC", "At the table"],
                    ["Elias Voss", "Core guide. Remembers fully. A person constructed by the Unremembered from stolen village memories, now independent and sincere. Starts the escort, first reset and termination argument."],
                    ["Mara Vale", "Core emotional anchor. Does not remember facts. Starts the Lantern House hospitality and fire; her grief and dreams make Ashfall worth saving."],
                    ["Nessa Grey", "Core recognition character. Remembers people and promises, not complete timelines. Confirms that something survived the reset."],
                    ["Mayor Aldren Morn", "Pressure character. Does not know the day repeats at the start of a loop. His recurring fear creates detention, control and preventable harm."],
                    ["Constable Iven Rook", "Investigation character. Forgets resets, but his surviving notebook turns theories into tests and exposes earlier expeditions."],
                    ["Silas Wren", "Mechanism character. Forgets facts but repeatedly rediscovers his guilt and diagrams. Explains how the crystal and wards work."],
                    ["Seraphine Vey", "Revelation character. Her fragmented echo remembers the catastrophe, denies creating Elias and explains the cost of each ending."],
                ],
                "key": "clues",
            },
            {"type": "h2", "text": "Story handoff"},
            {
                "type": "table",
                "rows": [
                    ["When the party needs...", "Use...", "The character starts..."],
                    ["A reason to enter", "Elias", "The escort and Crystal Chest."],
                    ["A reason to care", "Mara", "Hospitality, Tomas's empty place and the inn fire."],
                    ["Proof a relationship survived", "Nessa", "A remembered promise after the first reset."],
                    ["Testable evidence", "Iven", "A claim the next reset can prove."],
                    ["Immediate pressure", "Aldren", "A detention, curfew or harmful protective ritual."],
                    ["The mechanism", "Silas", "A clock disagreement, ward diagram or command sequence."],
                    ["The true cost", "Seraphine", "A warning through an anchor, then the final revelation."],
                ],
            },
            {"type": "h2", "text": "Voice shorthand"},
            {
                "type": "list",
                "items": [
                    "Elias: measured explanations, sincere concern, careful omissions and personal memories that occasionally belong to someone else.",
                    "Mara: warm questions, practical tasks and an extra place at the table.",
                    "Nessa: soft, literal statements about people and promises rather than cryptic lore.",
                    "Aldren: reassuring certainty that becomes defensive control.",
                    "Iven: short questions, written evidence and controlled suspicion.",
                    "Silas: rapid technical language interrupted by fear of the bell.",
                    "Seraphine: fragmented warnings at first; later severe honesty and exhausted compassion.",
                ],
            },
        ],
    )

    replace_section(
        document,
        "Elias Voss",
        [
            {"type": "p", "text": "The false survivor who genuinely loves Ashfall and unknowingly carries the Unremembered's escape plan."},
            {
                "type": "table",
                "rows": [
                    ["At a glance", "Roleplaying cues"],
                    ["Appears to be a travelling merchant and former Ashfall courier escorting a sealed relic. He remembers every reset and wants outsiders to restore his village to ordinary time.", "Speak calmly and helpfully. Avoid direct lies when omission works. Tell affectionate village stories with small details that later prove to belong to other people."],
                ],
                "key": "clues",
            },
            {"type": "h2", "text": "What drives them"},
            {
                "type": "table",
                "rows": [
                    ["Wants", "Fears"],
                    ["To stop the reset and let his people remember, age and build a future.", "That Seraphine intends to preserve Ashfall forever; that earlier expeditions died because he failed them; and that he may not be the man he remembers."],
                ],
                "key": "clues",
            },
            {"type": "h2", "text": "Hidden truth"},
            {
                "type": "callout",
                "text": "DM secret: The Unremembered assembled Elias from stolen memories as an unwitting route to lawful termination. The crystal anchored him into an independent person the entity cannot control. There was no original Elias Voss, but his present identity, love and choices are real.",
            },
            {"type": "h2", "text": "What they can reveal"},
            {
                "type": "list",
                "items": [
                    "The crystal anchors the party's memories and wakes them immediately before a reset.",
                    "He remembers growing up in Ashfall, escaping with the crystal and being sent back with living witnesses.",
                    "He has brought other groups and hidden the truth because earlier recruits refused or treated the villagers as disposable.",
                    "The final ward rejects him and prevents him from reading its deepest warning.",
                    "He believes termination leaves the villagers alive in ordinary time; he does not know their identities are the lock.",
                ],
            },
            {"type": "h2", "text": "How the party changes them"},
            {
                "type": "table",
                "rows": [
                    ["They move toward trust when...", "They move toward opposition when..."],
                    ["They treat constructed people as real, prove that his memories came from several villagers, distinguish his origin from his present choices and offer a future that also contains the entity.", "They call him a monster or puppet, dismiss the villagers as copies, preserve the loop without seeking another solution or accuse him of knowingly serving the entity."],
                ],
                "key": "clues",
            },
            {"type": "h2", "text": "Scenes they start"},
            {
                "type": "table",
                "rows": [
                    ["Trigger", "Elias's move", "Story result"],
                    ["Arrival", "Claims the chest must remain at the Lantern House until morning.", "Keeps the party inside Ashfall long enough to care about its people."],
                    ["3:12 a.m.", "Waits beside the humming chest while the crystal wakes the party.", "The reset becomes undeniable and his prior knowledge becomes suspicious."],
                    ["The party confronts his lie", "Admits the loop and crystal but initially conceals earlier expeditions.", "Moves the party toward Iven's notes, Mara's dreams and the manor records."],
                    ["His memories are disproved", "Rejects the evidence, then asks whether borrowed memories make his love meaningless.", "Opens the central personhood question and the possibility of changing his mind."],
                    ["Final chamber", "Stops outside the circle and argues for termination.", "Forces the party to answer his plan with evidence, relationship or rejection."],
                ],
                "key": "timeline",
            },
            {"type": "h2", "text": "Useful lines"},
            {"type": "p", "text": "\"You are remembering because the crystal was built to let witnesses learn before they choose.\""},
            {"type": "p", "text": "\"I lied about the road. I did not lie about wanting my people to live.\""},
            {"type": "p", "text": "\"If those memories belonged to someone else, why does losing them hurt me?\""},
            {"type": "p", "text": "\"Whatever made me, it does not get to choose what I do next.\""},
            {
                "type": "callout",
                "text": "Finale contribution: Convinced, Elias uses his connection to identify and oppose his creator. Unconvinced, he invokes termination because he still believes it will save Ashfall.",
            },
        ],
    )

    replacements = [
        (
            "A protector who has mistaken repetition for permission.",
            "A protector whose recurring fear turns protection into harm.",
        ),
        (
            "That the murders accomplish nothing; that the victims remember emotionally; and that he has become the danger he feared.",
            "That the ritual accomplishes nothing; that those he harms remember emotionally; and that he has become the danger he feared.",
        ),
        (
            "The party privately proves the village remains stable without a murder; protects his dignity while confronting his guilt.",
            "The party privately proves the village remains stable without his ritual and protects his dignity while confronting his guilt.",
        ),
        (
            "They humiliate him publicly, threaten the village or treat the murder as entertainment.",
            "They humiliate him publicly, threaten the village or treat the ritual's harm as entertainment.",
        ),
        (
            "I remember their faces. I simply refuse to call that memory proof.",
            "Every night I dream their faces before I choose. I refuse to call a dream proof.",
        ),
        (
            "A quiet nine-year-old who remembers every reset and speaks to the party as returning acquaintances.",
            "A quiet nine-year-old who retains people, promises and powerful moments across resets without remembering a complete timeline.",
        ),
        (
            "This lets her see across loops, but ending Ashfall may erase her first.",
            "This gives her partial impressions across loops, but ending Ashfall may erase her first.",
        ),
        (
            "Complication: Aldren’s ritual murder. Relationship: convince Aldren that repetition does not make cruelty harmless.",
            "Complication: Aldren detains someone he believes is a threat. Relationship: prove that recurring fear does not make repeated harm necessary.",
        ),
        (
            "Relationship: make Elias confront evidence that his merciful ending opens the prison.",
            "Relationship: make Elias confront evidence that his promised restoration opens the prison.",
        ),
        (
            "Elias always produces the same unexplained dread.",
            "Some of Elias's childhood stories feel painfully familiar to Mara because their details belonged to Tomas. She cannot explain why.",
        ),
        (
            "That Elias hid something from the \"lantern people.\"",
            "That Elias feels like several remembered people gathered into one person, although Nessa cannot explain why.",
        ),
        (
            "DM secret The Unremembered encouraged his ritual through dreams, but it does not control him. He chooses one victim each night because the victim returns. Fear and alienation create Fracture imprints useful to the entity.",
            "DM secret The Unremembered encourages Aldren's protective ritual through recurring dreams, but it does not control him. Aldren begins each loop believing the threatened villager could die permanently. If the party proves the reset during that day, he may rationalise harm because it will be undone; after reset he forgets the proof but retains emotional guilt or fear.",
        ),
        (
            "The Unremembered encouraged his ritual through dreams, but it does not control him. He chooses one victim each night because the victim returns. Fear and alienation create Fracture imprints useful to the entity.",
            "The Unremembered encourages Aldren's protective ritual through recurring dreams, but it does not control him. Aldren begins each loop believing the threatened villager could die permanently. If the party proves the reset during that day, he may rationalise harm because it will be undone; after reset he forgets the proof but retains emotional guilt or fear.",
        ),
        (
            "Evidence Elias has used more than one cover identity.",
            "Evidence Elias has used more than one cover identity and no record of an original Elias household.",
        ),
        (
            "DM secret Seraphine's physical body died completing the containment spell. What remains is a damaged consciousness distributed through its original anchors. She remembers the catastrophe and most loops, but can communicate only in fragments without diverting strength from the seal. She distrusts Elias's conclusion, yet recognises him as a real person acting from grief rather than possession.",
            "DM secret Seraphine's physical body died completing the containment spell. What remains is a damaged consciousness distributed through its original anchors. She remembers the catastrophe and most loops, but can communicate only in fragments without diverting strength from the seal. She did not create or send Elias. She recognises the Unremembered's signature in his origin and also recognises that the crystal has made him an independent real person rather than a controlled vessel.",
        ),
        (
            "DM secret Seraphine’s physical body died completing the containment spell. What remains is a damaged consciousness distributed through its original anchors. She remembers the catastrophe and most loops, but can communicate only in fragments without diverting strength from the seal. She distrusts Elias’s conclusion, yet recognises him as a real person acting from grief rather than possession.",
            "DM secret Seraphine's physical body died completing the containment spell. What remains is a damaged consciousness distributed through its original anchors. She remembers the catastrophe and most loops, but can communicate only in fragments without diverting strength from the seal. She did not create or send Elias. She recognises the Unremembered's signature in his origin and also recognises that the crystal has made him an independent real person rather than a controlled vessel.",
        ),
        (
            "Seraphine’s physical body died completing the containment spell. What remains is a damaged consciousness distributed through its original anchors. She remembers the catastrophe and most loops, but can communicate only in fragments without diverting strength from the seal. She distrusts Elias’s conclusion, yet recognises him as a real person acting from grief rather than possession.",
            "Seraphine's physical body died completing the containment spell. What remains is a damaged consciousness distributed through its original anchors. She remembers the catastrophe and most loops, but can communicate only in fragments without diverting strength from the seal. She did not create or send Elias. She recognises the Unremembered's signature in his origin and also recognises that the crystal has made him an independent real person rather than a controlled vessel.",
        ),
        (
            "An earlier party warned Elias that \"the door opens both ways,\" but he dismissed them.",
            "She never created or sent Elias; the Unremembered assembled him from stolen village memories, but the crystal now protects his independent identity.",
        ),
        (
            "An earlier party warned Elias that “the door opens both ways,” but he dismissed them.",
            "She never created or sent Elias; the Unremembered assembled him from stolen village memories, but the crystal now protects his independent identity.",
        ),
        (
            "Elias argues that no one should be imprisoned for another person's safety. Seraphine argues that freedom without containment is a calamity disguised as mercy. Both are partly right; the party's role is to discover an option neither could achieve alone.",
            "Elias argues that Ashfall can become an ordinary living village. Seraphine knows his expected restoration is a false memory designed to open the prison. She must still admit that his constructed origin does not make him or the villagers less real. The party's role is to discover an option neither could achieve alone.",
        ),
        (
            "Elias argues that no one should be imprisoned for another person’s safety. Seraphine argues that freedom without containment is a calamity disguised as mercy. Both are partly right; the party’s role is to discover an option neither could achieve alone.",
            "Elias argues that Ashfall can become an ordinary living village. Seraphine knows his expected restoration is a false memory designed to open the prison. She must still admit that his constructed origin does not make him or the villagers less real. The party's role is to discover an option neither could achieve alone.",
        ),
        (
            "that outsiders will accept Elias’s merciful but catastrophic termination",
            "that outsiders will accept Elias's hopeful but catastrophic restoration",
        ),
        (
            "A compassionate group trusted Elias and attempted to terminate the loop.",
            "A compassionate group trusted Elias's promise that termination would restore Ashfall to ordinary life and attempted to end the loop.",
        ),
        (
            "Elias's mercy - terminate Ashfall.",
            "Elias's restoration - terminate Ashfall.",
        ),
        (
            "Elias’s mercy — terminate Ashfall.",
            "Elias's restoration - terminate Ashfall.",
        ),
        (
            "Elias strongly opposes this ending unless convinced no safer alternative exists.",
            "Elias strongly opposes this ending unless convinced that his promised restoration was planted by the entity and that no immediate safer alternative exists.",
        ),
        (
            "Elias realises the truth only when it is too late.",
            "Elias realises that his promised restoration was the entity's escape plan only when it is too late.",
        ),
        (
            "Missing Truth risks repeating Elias's catastrophic version of release.",
            "Missing Truth risks repeating Elias's catastrophic restoration plan.",
        ),
        (
            "Missing Truth risks repeating Elias’s catastrophic version of release.",
            "Missing Truth risks repeating Elias's catastrophic restoration plan.",
        ),
        (
            "Occurs through Elias's uncorrected termination plan, a high-Fracture ritual failure or deliberate release.",
            "Occurs through Elias's uncorrected restoration plan, a high-Fracture ritual failure or deliberate release.",
        ),
        (
            "Occurs through Elias’s uncorrected termination plan, a high-Fracture ritual failure or deliberate release.",
            "Occurs through Elias's uncorrected restoration plan, a high-Fracture ritual failure or deliberate release.",
        ),
        (
            "Do not mistake Elias’s mercy for safety—or my safety for justice.",
            "Do not mistake Elias's hope for proof - or my safety for justice.",
        ),
        (
            "Who resists them? Aldren protects order; Iven demands proof; Elias insists that ending the loop is mercy; Silas avoids guilt; the entity offers an easier answer.",
            "Who resists them? Aldren protects order; Iven demands proof; Elias insists termination will restore Ashfall; Silas avoids guilt; the entity offers an easier answer.",
        ),
        (
            "Aldren protects order; Iven demands proof; Elias insists that ending the loop is mercy; Silas avoids guilt; the entity offers an easier answer.",
            "Aldren protects order; Iven demands proof; Elias insists termination will restore Ashfall; Silas avoids guilt; the entity offers an easier answer.",
        ),
        (
            "End a session on a contradiction: someone remembers, a grave changes, Elias appears in a portrait or the crystal uses a character's voice.",
            "End a session on a contradiction: someone remembers, a grave changes, Elias's childhood story belongs to someone else or the crystal uses a character's voice.",
        ),
        (
            "Elias brought earlier parties",
            "Elias brought earlier parties",
        ),
        (
            "Elias: convinced / doubtful / committed to termination",
            "Elias: origin revealed / convinced / committed to restoration",
        ),
    ]
    for old, new in replacements:
        replace_text_everywhere(document, old, new)

    for table in document.tables:
        if not table.rows or table.rows[0].cells[0].text.strip() != "Revelation":
            continue
        if any(row.cells[0].text.strip() == "Elias's constructed origin" for row in table.rows):
            continue
        row = table.add_row()
        set_cell_text(row.cells[0], "Elias's constructed origin")
        set_cell_text(
            row.cells[1],
            "□ The Unremembered created him from stolen memories, but the crystal made him independent.",
        )
        set_table_geometry(table, table_widths(2, "clues"))

    document.save(OUTPUT)
    print(f"Created {OUTPUT.name}")
    print(f"Paragraphs: {len(document.paragraphs)}")
    print(f"Tables: {len(document.tables)}")


if __name__ == "__main__":
    main()
